"""Generate an updated CV (DOCX) from portfolio data."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------- Styling helpers ----------
PRIMARY = RGBColor(0x0F, 0x4C, 0x75)   # deep blue
ACCENT = RGBColor(0x32, 0x99, 0xCC)    # cyan-blue
DARK = RGBColor(0x22, 0x2B, 0x36)
GREY = RGBColor(0x55, 0x6B, 0x7B)
LIGHT_GREY = RGBColor(0x90, 0xA4, 0xAE)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_line(paragraph, color=PRIMARY, size=12):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2]))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def run(p, text, *, bold=False, italic=False, size=10, color=DARK, font="Calibri"):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run(p, text.upper(), bold=True, size=12, color=PRIMARY, font="Calibri")
    add_horizontal_line(p)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size=10, color=DARK)
    return p


def kv_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, f"{label} : ", bold=True, size=10, color=PRIMARY)
    run(p, value, size=10, color=DARK)
    return p


# ---------- Build the document ----------
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ---------- Header (Name + Title) ----------
name_p = doc.add_paragraph()
name_p.paragraph_format.space_after = Pt(0)
run(name_p, "MAHASALOTRA Toandromananjara", bold=True, size=22, color=PRIMARY)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(4)
run(title_p, "Développeur Full-Stack JS/TS  |  AI Engineer", bold=True, size=12, color=ACCENT)

# Contact line
contact_p = doc.add_paragraph()
contact_p.paragraph_format.space_after = Pt(2)
run(contact_p, "mahasalotra@gmail.com", size=9, color=GREY)
run(contact_p, "   |   ", size=9, color=LIGHT_GREY)
run(contact_p, "+261 34 03 121 63", size=9, color=GREY)
run(contact_p, "   |   ", size=9, color=LIGHT_GREY)
run(contact_p, "Antananarivo, Madagascar", size=9, color=GREY)

links_p = doc.add_paragraph()
links_p.paragraph_format.space_after = Pt(2)
run(links_p, "Portfolio : toandro-portfolio.netlify.app", size=9, color=GREY)
run(links_p, "   |   ", size=9, color=LIGHT_GREY)
run(links_p, "LinkedIn : linkedin.com/in/toandromananjara-mahasalotra", size=9, color=GREY)
run(links_p, "   |   ", size=9, color=LIGHT_GREY)
run(links_p, "GitHub : github.com/ToandroMananjara", size=9, color=GREY)

# Separator line
sep = doc.add_paragraph()
sep.paragraph_format.space_after = Pt(2)
add_horizontal_line(sep, color=ACCENT, size=8)

# ---------- Profil ----------
section_heading(doc, "Profil professionnel")
profil_p = doc.add_paragraph()
profil_p.paragraph_format.space_after = Pt(4)
run(
    profil_p,
    "Développeur Full-Stack JavaScript/TypeScript et AI Engineer, je conçois des produits web et mobiles modernes — du frontend React/Next.js aux APIs NestJS/FastAPI — et j'intègre de plus en plus de fonctionnalités IA en production (LLMs, RAG, voix/TTS, automatisation). ",
    size=10, color=DARK,
)
run(
    profil_p,
    "Habitué aux environnements production et au travail collaboratif, j'aime livrer du code maintenable, performant et soigné — pensé pour la prochaine personne qui devra le maintenir.",
    size=10, color=DARK,
)

# ---------- Expertise ----------
section_heading(doc, "Domaines d'expertise")
expertises = [
    ("Frontend Engineering", "Interfaces React/Next.js pixel-perfect, accessibles et performantes. Composants réutilisables, animations soignées."),
    ("Backend & APIs", "Services NestJS / Node.js / FastAPI fiables — modèles de données limpides, authentification, déploiements prévisibles."),
    ("Mobile Development", "Apps React Native / Expo cross-platform iOS & Android, expérience cohérente, navigation fluide."),
    ("AI Engineering", "Features LLM qui passent en prod : retrieval (RAG), prompts calibrés, agents et expériences vocales (ElevenLabs, Trigger.dev)."),
]
for title, desc in expertises:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(3)
    run(p, "•  ", bold=True, size=10, color=ACCENT)
    run(p, f"{title} — ", bold=True, size=10, color=PRIMARY)
    run(p, desc, size=10, color=DARK)

# ---------- Expérience ----------
section_heading(doc, "Expérience professionnelle")

experiences_data = [
    {
        "title": "Développeur Full-Stack & AI Engineer",
        "company": "Redsmite",
        "type": "Freelance",
        "period": "2026 — Aujourd'hui",
        "location": "Madagascar / Remote",
        "bullets": [
            "Consultant sur des missions full-stack et IA pour différents clients.",
            "Intégration de fonctionnalités IA en production : LLMs, voix / TTS (ElevenLabs), workflows d'automatisation (Trigger.dev).",
            "Stack moderne : Next.js, TypeScript, FastAPI, Supabase — focus mise en production et fiabilité.",
        ],
    },
    {
        "title": "Développeur Full-Stack",
        "company": "Freelance",
        "type": "Freelance",
        "period": "Juin 2025 — Aujourd'hui",
        "location": "Remote",
        "bullets": [
            "Conception d'applications web complètes en Next.js, NestJS, TypeScript et PostgreSQL.",
            "Réécriture complète d'un backend en NestJS avec architecture modulaire et APIs REST.",
            "Refactorisation, amélioration de la qualité du code, nouvelles fonctionnalités frontend & backend.",
            "Modélisation et maintenance de bases PostgreSQL, déploiement VPS (Docker + Nginx).",
        ],
    },
    {
        "title": "Développeur Front-End",
        "company": "Freelance",
        "type": "Freelance",
        "period": "Août 2025",
        "location": "Remote",
        "bullets": [
            "Conception d'une application web innovante intégrant l'IA pour l'industrie musicale.",
            "Interface utilisateur dynamique et responsive en React et TypeScript.",
        ],
    },
    {
        "title": "Développeur Mobile",
        "company": "Paika Sarl",
        "type": "Stage",
        "period": "Juillet — Août 2025",
        "location": "Madagascar",
        "bullets": [
            "Développement d'une application mobile pour food trucks en React Native, NestJS, TypeScript et Zustand.",
            "Implémentation complète : catalogue produits, système de commandes, intégration paiement.",
            "Modélisation PostgreSQL, livraison de features en Agile Scrum, frontend/backend et tests.",
        ],
    },
    {
        "title": "Développeur Web",
        "company": "Kinga IT",
        "type": "Stage",
        "period": "Mars — Mai 2025",
        "location": "Madagascar",
        "bullets": [
            "Contribution full-stack en React, NestJS, TypeScript, MongoDB et PostgreSQL.",
            "Intégration de nouvelles fonctionnalités frontend & backend, en équipe Agile.",
            "Modélisation, requêtes et optimisation BDD ; bonnes pratiques Git, code review, CI/CD.",
        ],
    },
]

for exp in experiences_data:
    head = doc.add_paragraph()
    head.paragraph_format.space_before = Pt(6)
    head.paragraph_format.space_after = Pt(0)
    run(head, exp["title"], bold=True, size=11, color=DARK)
    run(head, f"  •  {exp['company']}", bold=True, size=11, color=ACCENT)
    run(head, f"  ({exp['type']})", italic=True, size=10, color=GREY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    run(sub, f"{exp['period']}  |  {exp['location']}", italic=True, size=9, color=GREY)

    for b in exp["bullets"]:
        bullet(doc, b)

# ---------- Projets ----------
section_heading(doc, "Projets sélectionnés")
projects_data = [
    ("Plateforme de Statistiques Football", "Plateforme full-stack agrégeant des stats foot en direct — frontend React typé, backend PHP MVC, MySQL.", "React, TypeScript, PHP, MVC, MySQL"),
    ("Varotra — App E-commerce mobile", "App mobile e-commerce cross-platform en React Native, Expo et TypeScript.", "React Native, Expo, TypeScript"),
    ("Watch Store", "Front-end e-commerce pour un catalogue de montres de luxe.", "React, CSS3"),
    ("Portfolio personnel", "Portfolio interactif multilingue (FR/EN) avec animations et particles background.", "React, Vite, Tailwind CSS, i18n"),
]
for title, desc, tech in projects_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, "•  ", bold=True, size=10, color=ACCENT)
    run(p, f"{title} — ", bold=True, size=10, color=DARK)
    run(p, desc, size=10, color=DARK)
    run(p, f"  [{tech}]", italic=True, size=9, color=GREY)

# ---------- Compétences techniques ----------
section_heading(doc, "Compétences techniques")
skills_data = [
    ("Frontend", "React.js, Next.js, React Native, TypeScript, Tailwind CSS, Material UI, Bootstrap, Redux Toolkit, Zustand, i18n"),
    ("Backend", "Node.js, NestJS, Express.js, FastAPI, APIs REST, Prisma, JWT"),
    ("Bases de données", "PostgreSQL, MySQL, MongoDB, Supabase"),
    ("IA & Automatisation", "OpenAI, Anthropic Claude, Gemini, LangChain, RAG / Vector DB, Prompt Engineering, ElevenLabs (TTS), Trigger.dev"),
    ("Langages", "JavaScript (ES6+), TypeScript, Python, Java, C, C++, PHP"),
    ("DevOps & Outils", "Git, GitHub, GitHub Actions, Docker, Nginx, CI/CD, Linux, VPS"),
    ("Tests / QA", "Jest"),
    ("Méthodologies", "Agile, Scrum, Code Review, Clean Code"),
]
for label, value in skills_data:
    kv_line(doc, label, value)

# ---------- Formation ----------
section_heading(doc, "Formation")
edu_data = [
    ("Master 2 — MISA (Informatique)", "Université d'Antananarivo  |  2025 — 2026 (en cours)"),
    ("Master 1 — MISA (Informatique)", "Université d'Antananarivo  |  2024 — 2025 (validé)"),
    ("Licence — MISA (Informatique)", "Université d'Antananarivo  |  2023 — 2024 (diplômé)"),
    ("Formation Full-Stack Web", "Digital Training Center  |  2023"),
    ("Baccalauréat Scientifique (Série C)", "Lycée Nanisana  |  2018"),
]
for title, sub in edu_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, "•  ", bold=True, size=10, color=ACCENT)
    run(p, title, bold=True, size=10, color=DARK)
    run(p, f"  —  {sub}", italic=True, size=9, color=GREY)

# ---------- Engagement ----------
section_heading(doc, "Engagement associatif")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run(p, "Responsable de la Performance", bold=True, size=10, color=DARK)
run(p, "  •  SMARTKAJY  ", bold=True, size=10, color=ACCENT)
run(p, "(Novembre 2024 — Aujourd'hui)", italic=True, size=9, color=GREY)
bullet(doc, "Analyse et optimisation des événements et activités techniques pour améliorer l'efficacité opérationnelle.")

# ---------- Langues ----------
section_heading(doc, "Langues")
langues = [
    ("Malagasy", "Langue maternelle"),
    ("Français", "Courant (B2)"),
    ("Anglais", "Technique / Lecture (B1)"),
]
for lang, level in langues:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, "•  ", bold=True, size=10, color=ACCENT)
    run(p, f"{lang} — ", bold=True, size=10, color=DARK)
    run(p, level, size=10, color=GREY)

# ---------- Save ----------
out_dir = "/Users/toandro/Documents/MyProject/portfolio-v2/src/assets/data"
out_path = os.path.join(out_dir, "CV_Toandro_2026.docx")
doc.save(out_path)
print(f"CV généré : {out_path}")
